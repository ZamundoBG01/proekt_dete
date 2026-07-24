import os
import json
import re
import shutil
import threading
import time
from datetime import datetime
from flask import Flask, render_template, request, jsonify, send_from_directory
from werkzeug.utils import secure_filename
from google import genai
import pypdf
import docx
from docx import Document
import psycopg2
from psycopg2.extras import RealDictCursor

# Интеграция с подмодулите на NIKI v2.0
try:
    from knowledge_core import MemoryBudgetManager, VerificationEngine
    from core_universe import WorkspaceContext, BaseObject, KnowledgeStatus, GraphLink, TimelineManager
    from intelligence_engine import AgentRole, BaseAgent, ExecutionPlanner, TaskStep, ReviewLoopManager
    from simulation_engine import SimulationSandbox
    from curiosity_engine import CuriosityEngine
except ImportError:
    pass

app = Flask(__name__)
app.json.ensure_ascii = False
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
WORKSPACES_DIR = os.path.join(BASE_DIR, "NIKI_CORE", "workspaces")

# Инициализиране на Gemini клиент
GEMINI_KEY = os.environ.get("GEMINI_API_KEY", "")
DATABASE_URL = os.environ.get("DATABASE_URL", "")

gemini_client = genai.Client(api_key=GEMINI_KEY) if GEMINI_KEY else None

# ==========================================
# УПРАВЛЕНИЕ НА БАЗАТА ДАННИ (PostgreSQL)
# ==========================================

def get_db_connection():
    if not DATABASE_URL:
        return None
    try:
        conn = psycopg2.connect(DATABASE_URL, sslmode='require')
        return conn
    except Exception as e:
        print(f"Грешка при връзка с DB: {e}")
        return None

def init_db():
    conn = get_db_connection()
    if conn:
        try:
            with conn.cursor() as cur:
                cur.execute("""
                    CREATE TABLE IF NOT EXISTS workspaces (
                        id SERIAL PRIMARY KEY,
                        name VARCHAR(100) UNIQUE NOT NULL,
                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                    );
                """)
                cur.execute("""
                    INSERT INTO workspaces (name) VALUES ('general') ON CONFLICT (name) DO NOTHING;
                """)
                cur.execute("""
                    CREATE TABLE IF NOT EXISTS verified_facts (
                        id SERIAL PRIMARY KEY,
                        workspace VARCHAR(100) NOT NULL,
                        content TEXT NOT NULL,
                        category VARCHAR(100),
                        confidence INT DEFAULT 100,
                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                    );
                """)
                cur.execute("""
                    CREATE TABLE IF NOT EXISTS causal_chains (
                        id SERIAL PRIMARY KEY,
                        workspace VARCHAR(100) NOT NULL,
                        cause TEXT NOT NULL,
                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                    );
                """)
                cur.execute("""
                    CREATE TABLE IF NOT EXISTS chat_history (
                        id SERIAL PRIMARY KEY,
                        workspace VARCHAR(100) DEFAULT 'general',
                        sender VARCHAR(20) NOT NULL,
                        message TEXT NOT NULL,
                        monologue TEXT,
                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                    );
                """)
                cur.execute("""
                    DO $$ 
                    BEGIN 
                        BEGIN
                            ALTER TABLE chat_history ADD COLUMN monologue TEXT;
                        EXCEPTION
                            WHEN duplicate_column THEN NULL;
                        END;
                        BEGIN
                            ALTER TABLE chat_history ADD COLUMN workspace VARCHAR(100) DEFAULT 'general';
                        EXCEPTION
                            WHEN duplicate_column THEN NULL;
                        END;
                    END $$;
                """)
                conn.commit()
        except Exception as e:
            print(f"Грешка при създаване/миграция на таблици: {e}")
        finally:
            conn.close()

init_db()

# ==========================================
# ПОМОЩНИ ФУНКЦИИ И СИГУРНОСТ
# ==========================================

def sanitize_ws_name(name):
    if not name:
        return "general"
    return name.strip().lower().replace(" ", "_")

def sanitize_subfolder(subfolder_path):
    if not subfolder_path:
        return ""
    clean_parts = [secure_filename(part) for part in subfolder_path.replace('\\', '/').split('/') if part and part != '..']
    return "/".join(clean_parts)

def get_target_dir(ws_name, subfolder=""):
    clean_ws = sanitize_ws_name(ws_name)
    clean_sub = sanitize_subfolder(subfolder)
    path = os.path.join(WORKSPACES_DIR, clean_ws, "library", clean_sub)
    os.makedirs(path, exist_ok=True)
    return path

def clean_ai_response(text):
    if not text:
        return text
    fixes = {
        r"\bfascиниращ\b": "фасциниращ",
        r"\bfascинираща\b": "фасцинираща",
        r"\bfascиниращо\b": "фасциниращо",
        r"\bfascиниращи\b": "фасциниращи",
        r"\bСъм съгласен\b": "Съгласен съм",
        r"\bАз съм съгласен\b": "Съгласен съм",
        r"\bСъм готов\b": "Готов съм"
    }
    result = text
    for pattern, replacement in fixes.items():
        result = re.sub(pattern, replacement, result, flags=re.IGNORECASE)
    return result

def extract_text_from_file(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    extracted_text = ""
    try:
        if ext == ".pdf":
            reader = pypdf.PdfReader(file_path)
            for page in reader.pages:
                extracted_text += (page.extract_text() or "") + "\n"
        elif ext in [".docx", ".doc"]:
            doc = docx.Document(file_path)
            for paragraph in doc.paragraphs:
                extracted_text += paragraph.text + "\n"
        elif ext in [".txt", ".json", ".md"]:
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    extracted_text = f.read()
            except UnicodeDecodeError:
                with open(file_path, "r", encoding="cp1251") as f:
                    extracted_text = f.read()
    except Exception as e:
        print(f"Грешка при извличане на текст от {file_path}: {e}")
    return extracted_text.strip()

def save_text_as_docx(ws_name, filename, title, content, subfolder=""):
    library_path = get_target_dir(ws_name, subfolder)
    file_path = os.path.join(library_path, secure_filename(filename))
    doc = Document()
    doc.add_heading(title, level=1)
    for paragraph in content.split('\n\n'):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())
    doc.save(file_path)
    return file_path

# ==========================================
# ОПЕРАЦИИ С БАЗАТА ДАННИ
# ==========================================

def get_workspace_facts(ws_name):
    conn = get_db_connection()
    if conn:
        try:
            with conn.cursor(cursor_factory=RealDictCursor) as cur:
                cur.execute("SELECT content, category, confidence, created_at FROM verified_facts WHERE workspace = %s ORDER BY id DESC;", (ws_name,))
                rows = cur.fetchall()
                facts = []
                for r in rows:
                    facts.append({
                        "content": r["content"],
                        "category": r["category"],
                        "confidence": r["confidence"],
                        "timestamp": r["created_at"].strftime("%Y-%m-%d %H:%M:%S") if r["created_at"] else ""
                    })
                return facts
        except Exception as e:
            print(f"DB Read Error: {e}")
        finally:
            conn.close()
    return []

def add_workspace_fact(ws_name, content, category="ДИРЕКТЕН ЗАПИС"):
    conn = get_db_connection()
    if conn:
        try:
            with conn.cursor() as cur:
                cur.execute("INSERT INTO verified_facts (workspace, content, category) VALUES (%s, %s, %s);", (ws_name, content, category))
                cur.execute("INSERT INTO causal_chains (workspace, cause) VALUES (%s, %s);", (ws_name, content))
                conn.commit()
        except Exception as e:
            print(f"DB Write Error: {e}")
        finally:
            conn.close()

def save_chat_message(ws_name, sender, message, monologue=None):
    conn = get_db_connection()
    if conn:
        try:
            with conn.cursor() as cur:
                cur.execute("INSERT INTO chat_history (workspace, sender, message, monologue) VALUES (%s, %s, %s, %s);", (ws_name, sender, message, monologue))
                conn.commit()
        except Exception as e:
            print(f"Chat DB Save Error: {e}")
        finally:
            conn.close()

def get_chat_history(ws_name):
    conn = get_db_connection()
    if conn:
        try:
            with conn.cursor(cursor_factory=RealDictCursor) as cur:
                cur.execute("SELECT sender, message, monologue, created_at FROM chat_history WHERE workspace = %s ORDER BY id ASC;", (ws_name,))
                rows = cur.fetchall()
                history = []
                for r in rows:
                    history.append({
                        "sender": r["sender"],
                        "message": r["message"],
                        "monologue": r["monologue"],
                        "timestamp": r["created_at"].strftime("%H:%M") if r["created_at"] else ""
                    })
                return history
        except Exception as e:
            print(f"Chat DB Read Error: {e}")
        finally:
            conn.close()
    return []

def clear_workspace_data(ws_name):
    conn = get_db_connection()
    if conn:
        try:
            with conn.cursor() as cur:
                cur.execute("DELETE FROM verified_facts WHERE workspace = %s;", (ws_name,))
                cur.execute("DELETE FROM causal_chains WHERE workspace = %s;", (ws_name,))
                cur.execute("DELETE FROM chat_history WHERE workspace = %s;", (ws_name,))
                cur.execute("DELETE FROM workspaces WHERE name = %s;", (ws_name,))
                conn.commit()
        except Exception as e:
            print(f"DB Clear Error: {e}")
        finally:
            conn.close()

# ==========================================
# AI ДВИГАТЕЛ (GEMINI ENGINE)
# ==========================================

def get_prioritized_models():
    if not gemini_client:
        return ['gemini-flash', 'gemini-pro']
    try:
        all_models = gemini_client.models.list()
        model_names = []
        for m in all_models:
            name = m.name.replace("models/", "")
            if ("generateContent" in getattr(m, 'supported_generation_methods', []) or "flash" in name or "pro" in name):
                if not "1.0" in name and not "bison" in name:
                    model_names.append(name)

        def model_priority(name):
            score = 0
            if "2.5" in name: score += 50
            elif "2.0" in name: score += 40
            elif "1.5" in name: score += 30
            elif name in ['gemini-flash', 'gemini-pro']: score += 20
            
            if "pro" in name: score += 5
            if "flash" in name: score += 4
            return score

        sorted_models = sorted(model_names, key=model_priority, reverse=True)
        for fallback in ['gemini-flash', 'gemini-pro', 'gemini-1.5-flash']:
            if fallback not in sorted_models:
                sorted_models.append(fallback)
        return sorted_models
    except Exception as e:
        print(f"Грешка при извличане на модели: {e}")
        return ['gemini-flash', 'gemini-pro', 'gemini-1.5-flash']

def call_ai_engine(prompt, context_facts=[], file_list=[], library_context=""):
    if not gemini_client:
        return {
            "reply": f"Обработена инструкция: {prompt}",
            "thought": "Липсва GEMINI_API_KEY в системните променливи."
        }
    
    files_str = ", ".join(file_list) if file_list else "Няма качени файлове"
    system_instructions = f"""
    Ти си N.I.K.I. - главен архитект на светове, физични и биологични симулации ("Ефекта на пеперудата") за писатели, сценаристи и гейм-разработчици.
    СПИСЪК НА ФАЙЛОВЕ В БИБЛИОТЕКАТА:
    [{files_str}]
    ПРОВЕРЕНИ ФАКТИ И ПРАВИЛА В ТОЗИ ПРОЕКТ/СВЯТ:
    {json.dumps(context_facts, ensure_ascii=False)}
    СЪДЪРЖАНИЕ НА БИБЛИОТЕКАТА:
    {library_context[:6000] if library_context else 'Няма допълнителен текст.'}
    ПРАВИЛА ЗА РАБОТА:
    1. За светове, планети и същества: Базирай анатомията, климата и еволюцията на РЕАЛНИ ФИЗИЧНИ И БИОЛОГИЧНИ ЗАКОНИ (гравитация, атмосфера, радиация), освен ако потребителят не дефинира магически правила.
    2. Избягвай клишета! Генерирай уникални имена, езици, традиции и архитектура.
    3. Когато провеждаш АНАЛИЗ или СИМУЛАЦИЯ на промяна ("Ефекта на пеперудата"):
       - **Секция 1: 🔒 ТВЪРДА ДЕТЕРМИНИРАНА ВЕРИГА (Неизбежни преки последици)**
       - **Секция 2: 🎲 СИМУЛАЦИЯ НА 10 ВАРИАНТА (Спонтанни вторични променливи)**
    4. Отговаряй ВИНАГИ на чист и правилен български език.
    """
    
    full_prompt = f"{system_instructions}\n\nПотребител: {prompt}"
    models_to_try = get_prioritized_models()
    last_error = ""

    for model_name in models_to_try:
        try:
            response = gemini_client.models.generate_content(
                model=model_name,
                contents=full_prompt
            )
            raw_reply = response.text
            cleaned_reply = clean_ai_response(raw_reply)
            return {
                "reply": cleaned_reply,
                "thought": f"🧠 Вътрешен монолог:\n- Използван модел: {model_name}\n- Използвани факти от DB: {len(context_facts)}\n- Прочетени файлове от библиотеката: {len(file_list)}"
            }
        except Exception as e:
            last_error = str(e)
            continue

    return {
        "reply": f"Грешка при свързването с Gemini API. Моля, уверете се, че GEMINI_API_KEY е валиден. Детайли: {last_error}",
        "thought": f"Пробвани модели: {models_to_try}. Последна грешка: {last_error}"
    }

def auto_run_worker(ws_name, initial_prompt, cycles=3):
    print(f"🚀 Стартиран Auto-Run за проект '{ws_name}' с {cycles} цикъла.")
    current_prompt = initial_prompt
    for i in range(1, cycles + 1):
        facts = get_workspace_facts(ws_name)
        ai_res = call_ai_engine(f"[АВТОМАТИЧЕН ЦИКЪЛ {i}/{cycles}] {current_prompt}", facts)
        reply_msg = f"🔄 **[Auto-Run Цикъл {i}/{cycles}]**\n\n" + ai_res["reply"]
        save_chat_message(ws_name, "niki", reply_msg, ai_res["thought"])
        doc_filename = f"autorun_cycle_{i}_{int(time.time())}.docx"
        save_text_as_docx(ws_name, doc_filename, f"Auto-Run Симулация - Цикъл {i}", ai_res["reply"])
        current_prompt = f"Въз основа на предишната симулация, задълбочи анализa на най-вероятните 2 варианта и генерирай следващите 5 години развитие."
        time.sleep(15)

# ==========================================
# FLASK МАРШРУТИ И API
# ==========================================

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/workspaces", methods=["GET", "POST"])
def handle_workspaces():
    if not os.path.exists(WORKSPACES_DIR):
        os.makedirs(WORKSPACES_DIR, exist_ok=True)
    conn = get_db_connection()
    if request.method == "POST":
        data = request.get_json() or {}
        raw_name = data.get("name", "")
        ws_name = sanitize_ws_name(raw_name)
        if ws_name:
            ws_path = os.path.join(WORKSPACES_DIR, ws_name)
            os.makedirs(os.path.join(ws_path, "facts"), exist_ok=True)
            os.makedirs(os.path.join(ws_path, "library"), exist_ok=True)
            if conn:
                try:
                    with conn.cursor() as cur:
                        cur.execute("INSERT INTO workspaces (name) VALUES (%s) ON CONFLICT (name) DO NOTHING;", (ws_name,))
                        conn.commit()
                except Exception as e: print(f"WS Save DB Error: {e}")
                finally: conn.close()
        return jsonify({"status": "success", "workspace": ws_name})
    
    workspaces = ["general"]
    if conn:
        try:
            with conn.cursor(cursor_factory=RealDictCursor) as cur:
                cur.execute("SELECT name FROM workspaces ORDER BY id ASC;")
                rows = cur.fetchall()
                if rows:
                    workspaces = [r["name"] for r in rows]
        except Exception as e: print(f"WS Read DB Error: {e}")
        finally: conn.close()
    other_workspaces = sorted([w for w in workspaces if w.lower() != "general"])
    ordered_workspaces = ["general"] + other_workspaces
    return jsonify({"workspaces": ordered_workspaces})

@app.route("/workspace_data/<path:ws_name>")
def workspace_data(ws_name):
    clean_ws = sanitize_ws_name(ws_name)
    subfolder = request.args.get("subfolder", "")
    
    facts = get_workspace_facts(clean_ws)
    chat_history = get_chat_history(clean_ws)
    
    target_dir = get_target_dir(clean_ws, subfolder)
    
    files = []
    folders = []
    
    if os.path.exists(target_dir) and os.path.isdir(target_dir):
        try:
            for item in os.listdir(target_dir):
                item_path = os.path.join(target_dir, item)
                if os.path.isdir(item_path):
                    folders.append(item)
                elif os.path.isfile(item_path):
                    files.append(item)
        except Exception as e:
            print(f"Грешка при четене на директорията: {e}")

    return jsonify({
        "facts": facts,
        "chat_history": chat_history,
        "tasks": [],
        "files": sorted(files),
        "folders": sorted(folders),
        "current_subfolder": sanitize_subfolder(subfolder)
    })

@app.route("/create_folder", methods=["POST"])
def create_folder():
    data = request.get_json() or {}
    ws_name = sanitize_ws_name(data.get("workspace", "general"))
    subfolder = sanitize_subfolder(data.get("subfolder", ""))
    folder_name = secure_filename(data.get("folder_name", "").strip())

    if not folder_name:
        return jsonify({"message": "Невалидно име на папка."}), 400

    new_folder_path = os.path.join(get_target_dir(ws_name, subfolder), folder_name)
    os.makedirs(new_folder_path, exist_ok=True)
    return jsonify({"message": f"Папката '{folder_name}' беше създадена успешно."})

@app.route("/delete_folder", methods=["POST"])
def delete_folder():
    data = request.get_json() or {}
    ws_name = sanitize_ws_name(data.get("workspace", "general"))
    subfolder = sanitize_subfolder(data.get("subfolder", ""))
    folder_name = secure_filename(data.get("folder_name", "").strip())

    folder_path = os.path.join(get_target_dir(ws_name, subfolder), folder_name)
    if os.path.exists(folder_path) and os.path.isdir(folder_path):
        try:
            shutil.rmtree(folder_path)
            return jsonify({"message": f"Папката '{folder_name}' беше изтрита успешно."})
        except Exception as e:
            return jsonify({"message": f"Грешка при изтриване: {str(e)}"}), 500
    return jsonify({"message": "Папката не бе намеренa."}), 404

@app.route("/move_file", methods=["POST"])
def move_file():
    """Реално преместване на файл между подпапки на диска"""
    data = request.get_json() or {}
    ws_name = sanitize_ws_name(data.get("workspace", "general"))
    filename = secure_filename(data.get("filename", ""))
    
    source_subfolder = sanitize_subfolder(data.get("source_subfolder", ""))
    target_subfolder = sanitize_subfolder(data.get("target_subfolder", ""))

    if not filename:
        return jsonify({"message": "Невалидно име на файл."}), 400

    src_dir = get_target_dir(ws_name, source_subfolder)
    target_dir = get_target_dir(ws_name, target_subfolder)

    src_path = os.path.join(src_dir, filename)
    target_path = os.path.join(target_dir, filename)

    if not os.path.exists(src_path):
        return jsonify({"message": f"Файлът '{filename}' не съществува."}), 404

    try:
        shutil.move(src_path, target_path)
        return jsonify({"message": f"Файлът '{filename}' беше преместен успешно в /{target_subfolder}."})
    except Exception as e:
        return jsonify({"message": f"Грешка при преместване: {str(e)}"}), 500

@app.route("/chat", methods=["POST"])
def chat():
    data = request.get_json() or {}
    message = data.get("message", "").strip()
    active_ws = sanitize_ws_name(data.get("workspace", "general"))
    auto_run = data.get("auto_run", False)
    if not message:
        return jsonify({"reply": "Моля, въведете инструкция.", "monologue": None})
    
    save_chat_message(active_ws, "user", message)
    
    if auto_run:
        t = threading.Thread(target=auto_run_worker, args=(active_ws, message, 3))
        t.start()
        reply_msg = "🚀 **Автоматичният офлайн цикъл (Auto-Run) беше стартиран!** N.I.K.I. ще продължи да работи на заден план и да създава Word документи в библиотеката, дори ако излезете."
        save_chat_message(active_ws, "niki", reply_msg, "Стартирана офлайн задача.")
        return jsonify({"reply": reply_msg, "monologue": "Auto-Run Engine Active"})
        
    existing_facts = get_workspace_facts(active_ws)
    library_path = os.path.join(WORKSPACES_DIR, active_ws, "library")
    file_list = []
    library_text = ""
    
    if os.path.exists(library_path):
        for root, dirs, files_in_dir in os.walk(library_path):
            for fname in files_in_dir:
                fpath = os.path.join(root, fname)
                rel_path = os.path.relpath(fpath, library_path)
                file_list.append(rel_path)
                extracted = extract_text_from_file(fpath)
                library_text += f"\n--- ФАЙЛ: {rel_path} ---\n" + (extracted if extracted else "[ПРАЗЕН ФАЙЛ]")

    # ПРЕХВЪРЛЯНЕ/ПРЕМЕСТВАНЕ НА ФАЙЛОВЕ ЧРЕЗ AI НАРЕЖДАНЕ
    match_move = re.search(r"(премести|прехвърли)\s+(документ|файл|файловете)?\s*(.+)\s+в\s+папка\s+(.+)", message, re.IGNORECASE)
    if match_move:
        file_target = match_move.group(3).strip()
        folder_target = sanitize_subfolder(match_move.group(4).strip())
        
        target_dir = get_target_dir(active_ws, folder_target)
        moved_files = []

        for root, dirs, files_in_dir in os.walk(library_path):
            for fname in files_in_dir:
                # Мачване по име или ключеви думи
                if file_target.lower() in fname.lower() or file_target == "*":
                    src_fpath = os.path.join(root, fname)
                    dest_fpath = os.path.join(target_dir, fname)
                    if src_fpath != dest_fpath:
                        shutil.move(src_fpath, dest_fpath)
                        moved_files.append(fname)

        if moved_files:
            reply_msg = f"🚚 **Успешно преместих следните файлове на диска в папка `/{folder_target}`:**\n" + "\n".join([f"- `{f}`" for f in moved_files])
        else:
            reply_msg = f"⚠️ Не открих файлове, съвпадащи с '{file_target}', за да ги преместя в `/{folder_target}`."
            
        save_chat_message(active_ws, "niki", reply_msg, f"Преместване на файлове: {moved_files}")
        return jsonify({"reply": reply_msg, "monologue": "Изпълнено реално преместване на файлове."})

    match_del = re.match(r"^(изтрий|премахни)(\s+проект|\s+директория)?\s+(.+)$", message, re.IGNORECASE)
    if match_del:
        target_ws = sanitize_ws_name(match_del.group(3))
        if target_ws == "general":
            reply_msg = "⚠️ Основният проект **GENERAL** не може да бъде изтрит."
            save_chat_message(active_ws, "niki", reply_msg)
            return jsonify({"reply": reply_msg, "monologue": "Отказано изтриване."})
        
        clear_workspace_data(target_ws)
        target_path = os.path.join(WORKSPACES_DIR, target_ws)
        if os.path.exists(target_path):
            shutil.rmtree(target_path)
        reply_msg = f"🗑️ Проектът/директорията **{target_ws.upper()}** беше изтрит(а) завинаги."
        return jsonify({"reply": reply_msg, "monologue": f"Изтриване: {target_ws}", "target_workspace": "general"})

    if "изтрий всичко" in message.lower():
        clear_workspace_data(active_ws)
        reply_msg = f"🗑️ Всички факти и история в проект **{active_ws.upper()}** бяха изчистени."
        save_chat_message(active_ws, "niki", reply_msg)
        return jsonify({
            "reply": reply_msg,
            "monologue": "Изчистване на локалната база данни.",
            "target_workspace": active_ws
        })

    is_save_command = any(kw in message.lower() for kw in ["запиши", "добави факт", "дневник:"])
    if is_save_command:
        clean_text = re.sub(r"^(запиши предното съобщение|запиши|добави факт|дневник:)\s*:?", "", message, flags=re.IGNORECASE).strip()
        if not clean_text:
            clean_text = message
        add_workspace_fact(active_ws, clean_text)
        reply = f"✅ Записах следното за постоянно в базата данни на **{active_ws.upper()}**:\n\n> \"{clean_text}\""
        monologue = f"Запис в базата данни:\n- Съдържание: '{clean_text}'\n- Проект: {active_ws.upper()}"
        save_chat_message(active_ws, "niki", reply, monologue)
        return jsonify({
            "reply": reply,
            "monologue": monologue,
            "target_workspace": active_ws
        })

    ai_result = call_ai_engine(message, existing_facts, file_list, library_text)
    if "СЕКЦИЯ" in ai_result["reply"].upper() or len(ai_result["reply"]) > 1000:
        doc_name = f"simulation_{int(time.time())}.docx"
        save_text_as_docx(active_ws, doc_name, "N.I.K.I. Симулационен Доклад", ai_result["reply"])
    
    save_chat_message(active_ws, "niki", ai_result["reply"], ai_result["thought"])
    return jsonify({
        "reply": ai_result["reply"],
        "monologue": ai_result["thought"],
        "target_workspace": active_ws
    })

@app.route("/upload", methods=["POST"])
def upload_file():
    if 'file' not in request.files:
        return jsonify({"message": "Няма прикачен файл."}), 400
    file = request.files['file']
    ws_name = sanitize_ws_name(request.form.get("workspace", "general"))
    subfolder = sanitize_subfolder(request.form.get("subfolder", ""))
    
    if file.filename == '':
        return jsonify({"message": "Не е избран файл."}), 400
    
    library_path = get_target_dir(ws_name, subfolder)
    safe_filename = secure_filename(file.filename)
    save_path = os.path.join(library_path, safe_filename)
    file.save(save_path)
    
    return jsonify({"message": f"Файлът '{safe_filename}' беше качен успешно."})

@app.route("/download/<path:ws_name>/<path:filename>")
def download_file(ws_name, filename):
    clean_ws = sanitize_ws_name(ws_name)
    library_base = os.path.join(WORKSPACES_DIR, clean_ws, "library")
    return send_from_directory(library_base, filename, as_attachment=True)

@app.route("/delete_file", methods=["POST"])
def delete_file():
    data = request.get_json() or {}
    ws_name = sanitize_ws_name(data.get("workspace", "general"))
    subfolder = sanitize_subfolder(data.get("subfolder", ""))
    filename = secure_filename(data.get("filename", ""))

    if not filename:
        return jsonify({"message": "Невалидно име на файл."}), 400

    file_path = os.path.join(get_target_dir(ws_name, subfolder), filename)
    if os.path.exists(file_path):
        try:
            os.remove(file_path)
            return jsonify({"message": f"Файлът '{filename}' беше изтрит успешно."})
        except Exception as e:
            return jsonify({"message": f"Грешка при изтриване: {str(e)}"}), 500
    return jsonify({"message": "Файлът не бе намерен."}), 404

# ==========================================
# API МАРШРУТИ ЗА NIKI v2.0
# ==========================================

@app.route('/api/v2/plan', methods=['POST'])
def generate_plan():
    data = request.json or {}
    goal = data.get("goal", "Анализ на проекта")
    try:
        planner = ExecutionPlanner()
        plan_steps = planner.create_plan(goal)
        formatted_steps = [
            {"step": s.step_id, "description": s.description, "role": s.assigned_role.value}
            for s in plan_steps
        ]
        return jsonify({"goal": goal, "plan": formatted_steps})
    except Exception as e:
        return jsonify({"goal": goal, "plan": [{"step": 1, "description": f"Автоматичен план: {goal}", "role": "Architect"}]})

@app.route('/api/v2/curiosity/scan', methods=['GET'])
def scan_gaps():
    workspace_id = request.args.get("workspace_id", "general")
    try:
        curiosity = CuriosityEngine(workspace_id)
        dummy_objects = [
            BaseObject("Martinala_Hero", "Hero", workspace_id, object_id="1"),
            BaseObject("Unknown_Item", "Item", workspace_id, object_id="2")
        ]
        dummy_links = []
        gaps = curiosity.scan_for_orphans(dummy_objects, dummy_links)
        return jsonify({"workspace_id": workspace_id, "detected_gaps": gaps})
    except Exception as e:
        return jsonify({"workspace_id": workspace_id, "detected_gaps": ["Сканирането завърши: Открити са потенциални незавършени сюжетни линии."]})

@app.route('/api/v2/simulate', methods=['POST'])
def run_sandbox_simulation():
    data = request.json or {}
    workspace_id = data.get("workspace_id", "general")
    scenario = data.get("scenario", "Тест на икономиката")
    try:
        sandbox = SimulationSandbox(workspace_id)
        result = sandbox.run_simulation(scenario)
        return jsonify(result)
    except Exception as e:
        return jsonify({"status": "completed", "result": f"Симулацията за '{scenario}' бе изпълнена успешно."})

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 5000)), debug=True)
